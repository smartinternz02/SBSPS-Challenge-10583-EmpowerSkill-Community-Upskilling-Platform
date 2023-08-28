from flask import Flask, render_template, request
import fitz
from docx import Document
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import base64
import ibm_db

app = Flask(__name__)

def connect_to_db():
    db_credentials = {
        "dsn": "DATABASE=R2ADMIN;HOSTNAME=localhost;PORT=25000;PROTOCOL=TCPIP;UID=db2admin;PWD=db2admin;",
        "username": "db2admin",
        "password": "db2admin"
    }
    return ibm_db.connect(db_credentials["dsn"], "", "")


def close_db_connection(conn):
    ibm_db.close(conn)
    
    #for creation of table
    #CREATE TABLE DB2ADMIN.PDF_FILES (FILENAME VARCHAR(255), WORD_COUNT INT, KEYWORD_PRESENCE INT)
    # 1 means yes

    # Insert operation
def insert_pdf_file(conn, filename, word_count, keyword_presence):
    sql = f"INSERT INTO pdf_files (filename, word_count, keyword_presence) VALUES ('{filename}', {word_count}, {keyword_presence})"
    stmt = ibm_db.exec_immediate(conn, sql)
    return ibm_db.num_rows(stmt)

def convert_pdf_to_word(pdf_data, keywords):
    pdf_document = fitz.open("pdf", pdf_data)
    doc = Document()

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text = page.get_text("text")

        # Add the extracted text to the Word document
        doc.add_paragraph(text)

    output = BytesIO()
    doc.save(output)
    pdf_document.close()
    return output

def count_words_in_pdf(pdf_data):
    pdf_document = fitz.open("pdf", pdf_data)
    total_words = 0

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text = page.get_text("text")
        words = text.split()
        total_words += len(words)

    pdf_document.close()
    return total_words

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        pdf_files = request.files.getlist("pdf_files")
        keywords = request.form["keywords"].split(",")
        results = []

        for pdf_file in pdf_files:
            if pdf_file and pdf_file.filename.endswith(".pdf"):
                pdf_data = pdf_file.read()

                word_count = count_words_in_pdf(pdf_data)
                word_content = convert_pdf_to_word(pdf_data, keywords)

                # Check if any of the keywords are present in the Word content
                word_text = "\n".join([para.text for para in Document(word_content).paragraphs])
                keyword_presence_list = [kw.lower() in word_text.lower() for kw in keywords]
                keyword_presence = any(keyword_presence_list)

                # Collect the keywords found in this PDF
                found_keywords = [kw for kw, found in zip(keywords, keyword_presence_list) if found]

                # Establish a database connection
                conn = connect_to_db()

                if keyword_presence:
                    # Insert the PDF details into the database
                    insert_pdf_file(conn, pdf_file.filename, word_count, keyword_presence)

                    results.append({
                        "filename": pdf_file.filename,
                        "word_count": word_count,
                        "word_content": word_content,
                        "thumbnail": generate_thumbnail(word_content),
                        "keyword_presence": keyword_presence,
                        "found_keywords": found_keywords
                    })

                # Close the database connection
                close_db_connection(conn)

        if keyword_presence:
            return render_template("result.html", results=results)
        else:
            return render_template("deselect.html")

    return render_template("index.html")

def generate_thumbnail(word_content):
    # Create a blank image for the thumbnail
    thumbnail = Image.new("RGB", (100, 100), "white")
    draw = ImageDraw.Draw(thumbnail)

    # Draw a placeholder text on the thumbnail
    font = ImageFont.load_default()
    text = "Thumbnail"
    text_bbox = draw.textbbox((0, 0), text, font=font)
    text_position = ((100 - (text_bbox[2] - text_bbox[0])) / 2, (100 - (text_bbox[3] - text_bbox[1])) / 2)
    draw.text(text_position, text, fill="black", font=font)

    # Convert the image to bytes and return
    thumbnail_bytes = BytesIO()
    thumbnail.save(thumbnail_bytes, format="PNG")
    return base64.b64encode(thumbnail_bytes.getvalue()).decode("utf-8")

if __name__ == "__main__":
    app.run(debug=True)
